# document_counter_app.py
import os
import json
import hashlib
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
import threading
import queue
import logging

# Web framework imports
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import pandas as pd

# Document processing imports
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import pytesseract
import cv2
import numpy as np

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Database setup
def init_db():
    """Initialize SQLite database for document tracking"""
    conn = sqlite3.connect('documents.db')
    cursor = conn.cursor()
    
    # Create tables
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_hash TEXT UNIQUE,
            file_size INTEGER,
            file_type TEXT,
            page_count INTEGER,
            word_count INTEGER,
            character_count INTEGER,
            upload_date TIMESTAMP,
            processed_date TIMESTAMP,
            status TEXT,
            metadata TEXT
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS batches (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            batch_name TEXT,
            created_date TIMESTAMP,
            total_documents INTEGER,
            total_pages INTEGER,
            total_words INTEGER
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS batch_documents (
            batch_id INTEGER,
            document_id INTEGER,
            FOREIGN KEY (batch_id) REFERENCES batches (id),
            FOREIGN KEY (document_id) REFERENCES documents (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Document processing classes
class DocumentCounter:
    """Base class for document counting"""
    
    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.filename = self.file_path.name
        self.file_size = self.file_path.stat().st_size
        self.file_type = self.file_path.suffix.lower()
        
    def get_file_hash(self):
        """Calculate file hash for duplicate detection"""
        hasher = hashlib.md5()
        with open(self.file_path, 'rb') as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        return hasher.hexdigest()
    
    def count_pages(self):
        """Count pages - to be overridden by child classes"""
        raise NotImplementedError
    
    def count_words(self):
        """Count words - to be overridden by child classes"""
        raise NotImplementedError
    
    def process(self):
        """Process document and return statistics"""
        try:
            pages = self.count_pages()
            words = self.count_words()
            chars = self.count_characters()
            
            return {
                'filename': self.filename,
                'file_path': str(self.file_path),
                'file_hash': self.get_file_hash(),
                'file_size': self.file_size,
                'file_type': self.file_type,
                'page_count': pages,
                'word_count': words,
                'character_count': chars,
                'status': 'success'
            }
        except Exception as e:
            logger.error(f"Error processing {self.filename}: {str(e)}")
            return {
                'filename': self.filename,
                'file_path': str(self.file_path),
                'status': 'failed',
                'error': str(e)
            }
    
    def count_characters(self):
        """Count characters - to be implemented by subclasses if needed"""
        return 0

class PDFCounter(DocumentCounter):
    def count_pages(self):
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return len(pdf_reader.pages)
    
    def count_words(self):
        word_count = 0
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                word_count += len(text.split())
        return word_count
    
    def count_characters(self):
        char_count = 0
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                char_count += len(text)
        return char_count

class WordCounter(DocumentCounter):
    def count_pages(self):
        doc = Document(self.file_path)
        # Approximate page count based on paragraphs
        return len(doc.paragraphs) // 25 + 1
    
    def count_words(self):
        doc = Document(self.file_path)
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text.split())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_count += len(cell.text.split())
        return word_count
    
    def count_characters(self):
        doc = Document(self.file_path)
        char_count = 0
        for paragraph in doc.paragraphs:
            char_count += len(paragraph.text)
        return char_count

class ExcelCounter(DocumentCounter):
    def count_pages(self):
        # Excel doesn't have pages, count sheets instead
        wb = openpyxl.load_workbook(self.file_path, read_only=True)
        return len(wb.sheetnames)
    
    def count_words(self):
        wb = openpyxl.load_workbook(self.file_path, read_only=True)
        word_count = 0
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell and isinstance(cell, str):
                        word_count += len(cell.split())
        return word_count

class ImageCounter(DocumentCounter):
    def count_pages(self):
        # Image files count as 1 page
        return 1
    
    def count_words(self):
        # Use OCR to extract text from images
        image = Image.open(self.file_path)
        text = pytesseract.image_to_string(image)
        return len(text.split())
    
    def count_characters(self):
        image = Image.open(self.file_path)
        text = pytesseract.image_to_string(image)
        return len(text)

class TextCounter(DocumentCounter):
    def count_pages(self):
        # Approximate pages based on lines (assuming 50 lines per page)
        with open(self.file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            return max(1, len(lines) // 50)
    
    def count_words(self):
        with open(self.file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            return len(text.split())
    
    def count_characters(self):
        with open(self.file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            return len(text)

def get_counter(file_path):
    """Factory function to get appropriate counter"""
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    
    counters = {
        '.pdf': PDFCounter,
        '.docx': WordCounter,
        '.xlsx': ExcelCounter,
        '.xls': ExcelCounter,
        '.jpg': ImageCounter,
        '.jpeg': ImageCounter,
        '.png': ImageCounter,
        '.tiff': ImageCounter,
        '.txt': TextCounter,
        '.md': TextCounter,
        '.csv': TextCounter
    }
    
    counter_class = counters.get(ext)
    if counter_class:
        return counter_class(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

# Database operations
class DocumentDB:
    @staticmethod
    def save_document(doc_data):
        """Save document info to database"""
        conn = sqlite3.connect('documents.db')
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO documents 
                (filename, file_path, file_hash, file_size, file_type, 
                 page_count, word_count, character_count, upload_date, 
                 processed_date, status, metadata)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc_data['filename'],
                doc_data['file_path'],
                doc_data.get('file_hash', ''),
                doc_data.get('file_size', 0),
                doc_data.get('file_type', ''),
                doc_data.get('page_count', 0),
                doc_data.get('word_count', 0),
                doc_data.get('character_count', 0),
                datetime.now(),
                datetime.now(),
                doc_data.get('status', 'pending'),
                json.dumps(doc_data.get('metadata', {}))
            ))
            
            doc_id = cursor.lastrowid
            conn.commit()
            return doc_id
        except Exception as e:
            logger.error(f"Database error: {str(e)}")
            conn.rollback()
            raise
        finally:
            conn.close()
    
    @staticmethod
    def get_all_documents():
        """Retrieve all documents"""
        conn = sqlite3.connect('documents.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, file_type, page_count, word_count, 
                   character_count, upload_date, status
            FROM documents
            ORDER BY upload_date DESC
        ''')
        
        docs = []
        for row in cursor.fetchall():
            docs.append({
                'id': row[0],
                'filename': row[1],
                'file_type': row[2],
                'page_count': row[3],
                'word_count': row[4],
                'character_count': row[5],
                'upload_date': row[6],
                'status': row[7]
            })
        
        conn.close()
        return docs
    
    @staticmethod
    def get_statistics():
        """Get overall statistics"""
        conn = sqlite3.connect('documents.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT 
                COUNT(*) as total_docs,
                SUM(page_count) as total_pages,
                SUM(word_count) as total_words,
                SUM(character_count) as total_chars,
                COUNT(CASE WHEN status = 'success' THEN 1 END) as processed,
                COUNT(CASE WHEN status = 'failed' THEN 1 END) as failed
            FROM documents
        ''')
        
        stats = cursor.fetchone()
        conn.close()
        
        return {
            'total_documents': stats[0] or 0,
            'total_pages': stats[1] or 0,
            'total_words': stats[2] or 0,
            'total_characters': stats[3] or 0,
            'processed_documents': stats[4] or 0,
            'failed_documents': stats[5] or 0
        }

# Web routes
@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_files():
    """Upload and process files"""
    if 'files' not in request.files:
        return jsonify({'error': 'No files uploaded'}), 400
    
    files = request.files.getlist('files')
    results = []
    
    # Create upload directory if it doesn't exist
    upload_dir = Path('uploads')
    upload_dir.mkdir(exist_ok=True)
    
    for file in files:
        if file.filename == '':
            continue
        
        # Save file
        file_path = upload_dir / file.filename
        file.save(file_path)
        
        try:
            # Process document
            counter = get_counter(file_path)
            doc_data = counter.process()
            
            # Save to database
            doc_id = DocumentDB.save_document(doc_data)
            doc_data['id'] = doc_id
            
            results.append(doc_data)
        except Exception as e:
            logger.error(f"Error processing {file.filename}: {str(e)}")
            results.append({
                'filename': file.filename,
                'status': 'failed',
                'error': str(e)
            })
    
    return jsonify({
        'message': 'Upload completed',
        'results': results
    })

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get all documents"""
    documents = DocumentDB.get_all_documents()
    return jsonify(documents)

@app.route('/api/statistics', methods=['GET'])
def get_statistics():
    """Get statistics"""
    stats = DocumentDB.get_statistics()
    return jsonify(stats)

@app.route('/api/export', methods=['GET'])
def export_data():
    """Export document data to CSV"""
    documents = DocumentDB.get_all_documents()
    
    # Convert to DataFrame
    df = pd.DataFrame(documents)
    
    # Save to CSV
    export_file = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    df.to_csv(export_file, index=False)
    
    return send_file(export_file, as_attachment=True)

# HTML Template
html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Counter Application</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1400px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            overflow: hidden;
        }
        
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }
        
        .header h1 {
            font-size: 2.5em;
            margin-bottom: 10px;
        }
        
        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            padding: 30px;
            background: #f8f9fa;
        }
        
        .stat-card {
            background: white;
            padding: 20px;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        
        .stat-number {
            font-size: 2.5em;
            font-weight: bold;
            color: #667eea;
        }
        
        .stat-label {
            color: #666;
            margin-top: 10px;
        }
        
        .upload-section {
            padding: 30px;
            border-bottom: 1px solid #e0e0e0;
        }
        
        .upload-area {
            border: 2px dashed #667eea;
            border-radius: 10px;
            padding: 40px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s;
        }
        
        .upload-area:hover {
            background: #f0f0ff;
            border-color: #764ba2;
        }
        
        .upload-area.drag-over {
            background: #e8eaf6;
            border-color: #764ba2;
        }
        
        .file-input {
            display: none;
        }
        
        .upload-btn {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 12px 30px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 16px;
            margin-top: 20px;
        }
        
        .documents-section {
            padding: 30px;
        }
        
        .documents-table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
        }
        
        .documents-table th,
        .documents-table td {
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #e0e0e0;
        }
        
        .documents-table th {
            background: #f8f9fa;
            font-weight: 600;
            color: #333;
        }
        
        .status-success {
            color: #4caf50;
            font-weight: bold;
        }
        
        .status-failed {
            color: #f44336;
            font-weight: bold;
        }
        
        .export-btn {
            background: #4caf50;
            color: white;
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            margin-bottom: 20px;
        }
        
        .loading {
            text-align: center;
            padding: 20px;
            display: none;
        }
        
        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 0 auto;
        }
        
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        @media (max-width: 768px) {
            .stats-grid {
                grid-template-columns: repeat(2, 1fr);
            }
            
            .documents-table {
                font-size: 12px;
            }
            
            .documents-table th,
            .documents-table td {
                padding: 8px;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📄 Document Counter Application</h1>
            <p>Upload and automatically count pages, words, and characters in your documents</p>
        </div>
        
        <div class="stats-grid" id="statsGrid">
            <div class="stat-card">
                <div class="stat-number" id="totalDocs">0</div>
                <div class="stat-label">Total Documents</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" id="totalPages">0</div>
                <div class="stat-label">Total Pages</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" id="totalWords">0</div>
                <div class="stat-label">Total Words</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" id="totalChars">0</div>
                <div class="stat-label">Total Characters</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" id="processedDocs">0</div>
                <div class="stat-label">Processed</div>
            </div>
            <div class="stat-card">
                <div class="stat-number" id="failedDocs">0</div>
                <div class="stat-label">Failed</div>
            </div>
        </div>
        
        <div class="upload-section">
            <div class="upload-area" id="uploadArea">
                <p>📁 Drag and drop files here or click to upload</p>
                <p style="font-size: 12px; margin-top: 10px;">Supported formats: PDF, DOCX, XLSX, JPG, PNG, TXT, CSV</p>
                <input type="file" id="fileInput" multiple class="file-input">
                <button class="upload-btn" onclick="document.getElementById('fileInput').click()">Choose Files</button>
            </div>
            <div class="loading" id="loading">
                <div class="spinner"></div>
                <p>Processing documents...</p>
            </div>
        </div>
        
        <div class="documents-section">
            <button class="export-btn" onclick="exportData()">📊 Export to CSV</button>
            <table class="documents-table">
                <thead>
                    <tr>
                        <th>Filename</th>
                        <th>Type</th>
                        <th>Pages</th>
                        <th>Words</th>
                        <th>Characters</th>
                        <th>Date</th>
                        <th>Status</th>
                    </tr>
                </thead>
                <tbody id="documentsList">
                    <tr>
                        <td colspan="7" style="text-align: center;">No documents uploaded yet</td>
                    </tr>
                </tbody>
            </table>
        </div>
    </div>
    
    <script>
        // Load initial data
        loadDocuments();
        loadStatistics();
        
        // Upload area drag and drop
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        
        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('drag-over');
        });
        
        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('drag-over');
        });
        
        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('drag-over');
            const files = Array.from(e.dataTransfer.files);
            uploadFiles(files);
        });
        
        fileInput.addEventListener('change', (e) => {
            const files = Array.from(e.target.files);
            uploadFiles(files);
        });
        
        async function uploadFiles(files) {
            const formData = new FormData();
            files.forEach(file => {
                formData.append('files', file);
            });
            
            document.getElementById('loading').style.display = 'block';
            
            try {
                const response = await fetch('/api/upload', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                console.log('Upload result:', result);
                
                // Refresh data
                loadDocuments();
                loadStatistics();
            } catch (error) {
                console.error('Upload error:', error);
                alert('Error uploading files: ' + error.message);
            } finally {
                document.getElementById('loading').style.display = 'none';
                fileInput.value = '';
            }
        }
        
        async function loadDocuments() {
            try {
                const response = await fetch('/api/documents');
                const documents = await response.json();
                
                const tbody = document.getElementById('documentsList');
                if (documents.length === 0) {
                    tbody.innerHTML = '<tr><td colspan="7" style="text-align: center;">No documents uploaded yet</td></tr>';
                    return;
                }
                
                tbody.innerHTML = documents.map(doc => `
                    <tr>
                        <td>${escapeHtml(doc.filename)}</td>
                        <td>${doc.file_type || '-'}</td>
                        <td>${doc.page_count || 0}</td>
                        <td>${doc.word_count || 0}</td>
                        <td>${doc.character_count || 0}</td>
                        <td>${new Date(doc.upload_date).toLocaleDateString()}</td>
                        <td class="status-${doc.status}">${doc.status}</td>
                    </tr>
                `).join('');
            } catch (error) {
                console.error('Error loading documents:', error);
            }
        }
        
        async function loadStatistics() {
            try {
                const response = await fetch('/api/statistics');
                const stats = await response.json();
                
                document.getElementById('totalDocs').textContent = stats.total_documents || 0;
                document.getElementById('totalPages').textContent = stats.total_pages || 0;
                document.getElementById('totalWords').textContent = stats.total_words || 0;
                document.getElementById('totalChars').textContent = stats.total_characters || 0;
                document.getElementById('processedDocs').textContent = stats.processed_documents || 0;
                document.getElementById('failedDocs').textContent = stats.failed_documents || 0;
            } catch (error) {
                console.error('Error loading statistics:', error);
            }
        }
        
        async function exportData() {
            try {
                window.location.href = '/api/export';
            } catch (error) {
                console.error('Export error:', error);
                alert('Error exporting data');
            }
        }
        
        function escapeHtml(text) {
            const div = document.createElement('div');
            div.textContent = text;
            return div.innerHTML;
        }
        
        // Auto-refresh every 30 seconds
        setInterval(() => {
            loadDocuments();
            loadStatistics();
        }, 30000);
    </script>
</body>
</html>
"""

# Save HTML template
def save_template():
    """Save HTML template to templates directory"""
    template_dir = Path('templates')
    template_dir.mkdir(exist_ok=True)
    
    with open(template_dir / 'index.html', 'w') as f:
        f.write(html_template)

# Requirements file
requirements = """
flask==2.3.2
flask-cors==4.0.0
pandas==2.0.3
PyPDF2==3.0.1
python-docx==0.8.11
openpyxl==3.1.2
Pillow==10.0.0
pytesseract==0.3.10
opencv-python==4.8.0.74
numpy==1.24.3
"""

# Main application
if __name__ == '__main__':
    # Initialize database
    init_db()
    
    # Save template
    save_template()
    
    # Create upload directory
    Path('uploads').mkdir(exist_ok=True)
    
    # Run application
    app.run(debug=True, host='0.0.0.0', port=5000)